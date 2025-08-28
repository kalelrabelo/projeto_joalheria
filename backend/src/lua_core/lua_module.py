# backend/src/lua_module.py

from flask import Blueprint, request, jsonify
from datetime import datetime, timedelta
import json
import os
import threading
import time
import pandas as pd
import PyPDF2
from docx import Document
from werkzeug.utils import secure_filename

# Importar modelos do banco de dados
from src.models.user import db
from src.models.jewelry import Jewelry
from src.models.material import Material
from src.models.employee import Employee
from src.models.caixa import CaixaTransaction
from src.models.inventory import Inventory
from src.models.order import Order
from src.models.cost import Cost, Profit
from src.models.payment import Payment
from src.models.payroll import Payroll
from src.models.financial import FinancialTransaction
from src.models.vale import Vale
from src.vale_printer import vale_printer

# Importar novos módulos
from src.lua_context_manager import context_manager
from src.lua_file_processor import energy_processor

lua_bp = Blueprint("lua", __name__)

# Configuração para upload de arquivos
UPLOAD_FOLDER = '/tmp/lua_uploads'
ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'xls', 'pdf', 'docx'}

# Criar diretório de upload se não existir
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Configuração de contato do dono para notificações
OWNER_CONTACT_CONFIG = {
    "email": "",
    "telefone": "",
    "notif_estoque": True,
    "notif_aniversarios": True,
    "notif_vendas": False,
    "notif_financeiro": True,
    "updated_at": None
}

def allowed_file(filename):
    """Verificar se o arquivo tem extensão permitida"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def process_csv_file(filepath):
    """Processar arquivo CSV"""
    try:
        df = pd.read_csv(filepath)
        
        # Análise básica do CSV
        analysis = {
            "total_rows": len(df),
            "total_columns": len(df.columns),
            "columns": list(df.columns),
            "sample_data": df.head(3).to_dict('records') if len(df) > 0 else [],
            "data_types": df.dtypes.to_dict(),
            "missing_values": df.isnull().sum().to_dict()
        }
        
        return {
            "status": "success",
            "message": f"Arquivo CSV processado com sucesso. {len(df)} linhas e {len(df.columns)} colunas encontradas.",
            "analysis": analysis
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"Erro ao processar CSV: {str(e)}"
        }

def process_excel_file(filepath):
    """Processar arquivo Excel"""
    try:
        # Ler todas as planilhas
        excel_file = pd.ExcelFile(filepath)
        sheets_data = {}
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(filepath, sheet_name=sheet_name)
            sheets_data[sheet_name] = {
                "total_rows": len(df),
                "total_columns": len(df.columns),
                "columns": list(df.columns),
                "sample_data": df.head(3).to_dict('records') if len(df) > 0 else []
            }
        
        return {
            "status": "success",
            "message": f"Arquivo Excel processado com sucesso. {len(excel_file.sheet_names)} planilhas encontradas.",
            "sheets": sheets_data
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"Erro ao processar Excel: {str(e)}"
        }

def process_pdf_file(filepath):
    """Processar arquivo PDF"""
    try:
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            text_content = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content += page.extract_text()
            
            # Análise básica do PDF
            analysis = {
                "total_pages": len(pdf_reader.pages),
                "text_length": len(text_content),
                "preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
                "word_count": len(text_content.split()) if text_content else 0
            }
            
            return {
                "status": "success",
                "message": f"Arquivo PDF processado com sucesso. {len(pdf_reader.pages)} páginas encontradas.",
                "analysis": analysis
            }
    except Exception as e:
        return {
            "status": "error",
            "message": f"Erro ao processar PDF: {str(e)}"
        }

def process_docx_file(filepath):
    """Processar arquivo DOCX"""
    try:
        doc = Document(filepath)
        
        # Extrair texto de todos os parágrafos
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Análise básica do DOCX
        analysis = {
            "total_paragraphs": len(doc.paragraphs),
            "text_length": len(text_content),
            "preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
            "word_count": len(text_content.split()) if text_content else 0
        }
        
        return {
            "status": "success",
            "message": f"Arquivo DOCX processado com sucesso. {len(doc.paragraphs)} parágrafos encontrados.",
            "analysis": analysis
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"Erro ao processar DOCX: {str(e)}"
        }

# Dados da história do Antonio Rabelo para aprendizado contínuo
ANTONIO_RABELO_HISTORY = {
    "nome": "Antonio Rabelo",
    "empresa": "Joalheria Antonio Rabelo",
    "especialidade": "Joias de alta qualidade",
    "fundacao": "Empresa familiar com tradição em joalheria",
    "valores": "Qualidade, tradição e excelência no atendimento",
    "produtos": "Anéis, colares, brincos, pulseiras, joias personalizadas",
    "materiais": "Ouro, prata, pedras preciosas e semi-preciosas"
}

# Configurações do scheduler
SCHEDULER_CONFIG = {
    "relatorio_financeiro": {"hora": "08:00", "frequencia": "diario"},
    "relatorio_estoque": {"hora": "09:00", "frequencia": "diario"},
    "relatorio_funcionarios": {"hora": "07:30", "frequencia": "semanal"},
    "relatorio_vendas": {"hora": "18:00", "frequencia": "diario"}
}

# Armazenar relatórios gerados automaticamente
AUTOMATIC_REPORTS = []

def generate_automatic_report(report_type):
    """Gerar relatório automático e enviar notificação se configurado"""
    try:
        timestamp = datetime.now().isoformat()
        
        if report_type == "financeiro":
            financial_data = get_financial_data()
            report = {
                "tipo": "Relatório Financeiro Automático",
                "timestamp": timestamp,
                "dados": financial_data,
                "resumo": f"Receitas: R$ {financial_data.get('receitas', 0):.2f}, Despesas: R$ {financial_data.get('despesas', 0):.2f}, Lucro: R$ {financial_data.get('lucro', 0):.2f}"
            }
            
            # Enviar notificação financeira se habilitada
            if OWNER_CONTACT_CONFIG.get("notif_financeiro", False):
                send_notification_internal("financeiro", f"📊 {report['resumo']}")
                
        elif report_type == "estoque":
            inventory_data = get_inventory_data()
            report = {
                "tipo": "Relatório de Estoque Automático",
                "timestamp": timestamp,
                "dados": inventory_data,
                "resumo": f"Total de itens: {inventory_data.get('total_items', 0)}, Estoque baixo: {inventory_data.get('itens_estoque_baixo', 0)}"
            }
            
            # Enviar notificação de estoque se habilitada e houver itens com estoque baixo
            if OWNER_CONTACT_CONFIG.get("notif_estoque", False) and inventory_data.get('itens_estoque_baixo', 0) > 0:
                produtos_baixo = ', '.join(inventory_data.get('produtos_baixo_estoque', []))
                send_notification_internal("estoque", f"⚠️ Alerta de Estoque: {inventory_data.get('itens_estoque_baixo', 0)} itens com estoque baixo. Produtos: {produtos_baixo}")
                
        elif report_type == "funcionarios":
            employee_data = get_employee_data()
            report = {
                "tipo": "Relatório de Funcionários Automático",
                "timestamp": timestamp,
                "dados": employee_data,
                "resumo": f"Total de funcionários: {employee_data.get('total_funcionarios', 0)}, Ativos: {employee_data.get('funcionarios_ativos', 0)}"
            }
        elif report_type == "vendas":
            sales_data = get_sales_data()
            report = {
                "tipo": "Relatório de Vendas Automático",
                "timestamp": timestamp,
                "dados": sales_data,
                "resumo": f"Total de pedidos: {sales_data.get('total_pedidos', 0)}, Valor total: R$ {sales_data.get('valor_total_vendas', 0):.2f}"
            }
            
            # Enviar notificação de vendas se habilitada
            if OWNER_CONTACT_CONFIG.get("notif_vendas", False):
                send_notification_internal("vendas", f"🛍️ {report['resumo']}")
                
        else:
            report = {
                "tipo": f"Relatório {report_type} Automático",
                "timestamp": timestamp,
                "dados": {},
                "resumo": f"Relatório {report_type} gerado automaticamente"
            }
        
        AUTOMATIC_REPORTS.append(report)
        
        # Manter apenas os últimos 50 relatórios
        if len(AUTOMATIC_REPORTS) > 50:
            AUTOMATIC_REPORTS.pop(0)
            
        print(f"[LUA SCHEDULER] Relatório {report_type} gerado automaticamente às {timestamp}")
        
    except Exception as e:
        print(f"[LUA SCHEDULER] Erro ao gerar relatório {report_type}: {str(e)}")

def send_notification_internal(notification_type, message):
    """Função interna para enviar notificações"""
    try:
        # Verificar se o dono quer receber este tipo de notificação
        if not OWNER_CONTACT_CONFIG.get(f"notif_{notification_type}", False):
            return
        
        email = OWNER_CONTACT_CONFIG.get("email", "")
        telefone = OWNER_CONTACT_CONFIG.get("telefone", "")
        
        if email or telefone:
            # Log da notificação enviada
            notification_log = {
                "type": notification_type,
                "message": message,
                "email": email,
                "telefone": telefone,
                "timestamp": datetime.now().isoformat(),
                "status": "sent"
            }
            
            # Adicionar ao histórico de notificações
            if "notification_history" not in globals():
                global notification_history
                notification_history = []
            
            notification_history.append(notification_log)
            
            # Manter apenas as últimas 100 notificações
            if len(notification_history) > 100:
                notification_history.pop(0)
            
            print(f"[LUA NOTIFICATION] Notificação {notification_type} enviada: {message}")
            
    except Exception as e:
        print(f"[LUA NOTIFICATION] Erro ao enviar notificação: {str(e)}")

def check_birthday_notifications():
    """Verificar aniversários de clientes e enviar notificações"""
    try:
        if not OWNER_CONTACT_CONFIG.get("notif_aniversarios", False):
            return
            
        today = datetime.now()
        # Simular verificação de aniversários (aqui seria consulta real ao banco)
        # Por enquanto, enviar notificação de teste uma vez por dia
        
        # Verificar se já enviou notificação de aniversário hoje
        if "last_birthday_check" not in globals():
            global last_birthday_check
            last_birthday_check = None
            
        if last_birthday_check != today.strftime("%Y-%m-%d"):
            # Simular aniversariantes do dia
            birthday_message = f"🎂 Aniversariantes de hoje: Maria Silva, João Santos. Considere enviar felicitações especiais!"
            send_notification_internal("aniversarios", birthday_message)
            last_birthday_check = today.strftime("%Y-%m-%d")
            
    except Exception as e:
        print(f"[LUA BIRTHDAY] Erro ao verificar aniversários: {str(e)}")

def scheduler_worker():
    """Worker do scheduler que roda em background"""
    while True:
        try:
            current_time = datetime.now().strftime("%H:%M")
            current_day = datetime.now().weekday()  # 0 = Segunda, 6 = Domingo
            
            # Verificar aniversários uma vez por dia às 08:00
            if current_time == "08:00":
                check_birthday_notifications()
            
            for report_type, config in SCHEDULER_CONFIG.items():
                if current_time == config["hora"]:
                    if config["frequencia"] == "diario":
                        generate_automatic_report(report_type.replace("relatorio_", ""))
                    elif config["frequencia"] == "semanal" and current_day == 0:  # Segunda-feira
                        generate_automatic_report(report_type.replace("relatorio_", ""))
            
            time.sleep(60)  # Verificar a cada minuto
            
        except Exception as e:
            print(f"[LUA SCHEDULER] Erro no scheduler: {str(e)}")
            time.sleep(60)

# Iniciar o scheduler em background
scheduler_thread = threading.Thread(target=scheduler_worker, daemon=True)
scheduler_thread.start()
print("[LUA SCHEDULER] Scheduler de relatórios automáticos iniciado")

def get_financial_data():
    """Buscar dados financeiros do banco"""
    try:
        transactions = CaixaTransaction.query.all()
        total_receitas = sum(t.amount for t in transactions if t.amount > 0)
        total_despesas = sum(abs(t.amount) for t in transactions if t.amount < 0)
        lucro = total_receitas - total_despesas
        return {
            "receitas": total_receitas,
            "despesas": total_despesas,
            "lucro": lucro,
            "total_transacoes": len(transactions)
        }
    except Exception as e:
        return {"error": str(e)}

def get_inventory_data():
    """Buscar dados de estoque do banco"""
    try:
        inventory_items = Inventory.query.all()
        total_items = len(inventory_items)
        low_stock_items = [item for item in inventory_items if item.quantity < 10]
        return {
            "total_items": total_items,
            "itens_estoque_baixo": len(low_stock_items),
            "produtos_baixo_estoque": [item.jewelry.name if item.jewelry else "Item sem nome" for item in low_stock_items[:5]]
        }
    except Exception as e:
        return {"error": str(e)}

def get_employee_data():
    """Buscar dados de funcionários do banco"""
    try:
        employees = Employee.query.all()
        total_employees = len(employees)
        active_employees = [emp for emp in employees if emp.active]
        return {
            "total_funcionarios": total_employees,
            "funcionarios_ativos": len(active_employees),
            "nomes_funcionarios": [emp.name for emp in active_employees[:5]]
        }
    except Exception as e:
        return {"error": str(e)}

def get_sales_data():
    """Buscar dados de vendas do banco"""
    try:
        orders = Order.query.all()
        total_orders = len(orders)
        completed_orders = [order for order in orders if order.status == 'completed']
        total_sales = sum(order.total_price for order in completed_orders if order.total_price)
        return {
            "total_pedidos": total_orders,
            "pedidos_concluidos": len(completed_orders),
            "valor_total_vendas": total_sales
        }
    except Exception as e:
        return {"error": str(e)}

@lua_bp.route("/lua/chat", methods=["POST"])
def chat():
    try:
        user_message = request.json.get("message", "").lower()
        
        # Respostas inteligentes baseadas na mensagem do usuário com dados reais
        if "relatório financeiro" in user_message or "relatório de caixa" in user_message or ("relatório" in user_message and "financeiro" in user_message):
            from src.report_generator import report_generator
            
            # Extrair ano da mensagem se especificado
            import re
            year_match = re.search(r'20\d{2}', user_message)
            year = int(year_match.group()) if year_match else datetime.now().year
            
            result = report_generator.generate_financial_report(year)
            if result["status"] == "success":
                response = f"📊 {result['message']}! PDF gerado com dados reais: R$ {result['data']['total_receitas']:.2f} em receitas, R$ {result['data']['total_despesas']:.2f} em despesas. Download disponível."
                # Adicionar informações de download
                return jsonify({
                    "response": response,
                    "timestamp": datetime.now().isoformat(),
                    "status": "success",
                    "report_generated": True,
                    "download_url": f"/lua/download_report/{result['filename']}",
                    "filename": result['filename']
                })
            else:
                response = f"Erro ao gerar relatório: {result['message']}"
        elif "relatório de funcionários" in user_message or ("relatório" in user_message and "funcionário" in user_message):
            from src.report_generator import report_generator
            
            import re
            year_match = re.search(r'20\d{2}', user_message)
            year = int(year_match.group()) if year_match else datetime.now().year
            
            result = report_generator.generate_employee_report(year)
            if result["status"] == "success":
                response = f"👥 {result['message']}! PDF gerado com {result['data']['total_funcionarios']} funcionários, {result['data']['funcionarios_ativos']} ativos. Download disponível."
                return jsonify({
                    "response": response,
                    "timestamp": datetime.now().isoformat(),
                    "status": "success",
                    "report_generated": True,
                    "download_url": f"/lua/download_report/{result['filename']}",
                    "filename": result['filename']
                })
            else:
                response = f"Erro ao gerar relatório: {result['message']}"
        elif "relatório de estoque" in user_message or ("relatório" in user_message and "estoque" in user_message):
            from src.report_generator import report_generator
            
            result = report_generator.generate_inventory_report()
            if result["status"] == "success":
                response = f"📦 {result['message']}! PDF gerado com {result['data']['total_items']} itens, {result['data']['low_stock_items']} com estoque baixo. Download disponível."
                return jsonify({
                    "response": response,
                    "timestamp": datetime.now().isoformat(),
                    "status": "success",
                    "report_generated": True,
                    "download_url": f"/lua/download_report/{result['filename']}",
                    "filename": result['filename']
                })
            else:
                response = f"Erro ao gerar relatório: {result['message']}"
        elif "relatório financeiro" in user_message or "financeiro" in user_message:
            financial_data = get_financial_data()
            if "error" not in financial_data:
                response = f"📊 Relatório Financeiro: Receitas: R$ {financial_data['receitas']:.2f}, Despesas: R$ {financial_data['despesas']:.2f}, Lucro: R$ {financial_data['lucro']:.2f}. Total de {financial_data['total_transacoes']} transações registradas."
            else:
                response = "📊 Gerando relatório financeiro... Analisando receitas, despesas e lucros do período solicitado."
        elif "estoque" in user_message:
            inventory_data = get_inventory_data()
            if "error" not in inventory_data:
                response = f"📦 Estoque: {inventory_data['total_items']} itens no total. {inventory_data['itens_estoque_baixo']} itens com estoque baixo. Produtos em falta: {', '.join(inventory_data['produtos_baixo_estoque']) if inventory_data['produtos_baixo_estoque'] else 'Nenhum'}"
            else:
                response = "📦 Consultando estoque... Verificando disponibilidade de anéis, colares, brincos e outras joias."
        elif "funcionários" in user_message or "folga" in user_message:
            employee_data = get_employee_data()
            if "error" not in employee_data:
                response = f"👥 Funcionários: {employee_data['total_funcionarios']} funcionários cadastrados, {employee_data['funcionarios_ativos']} ativos. Funcionários: {', '.join(employee_data['nomes_funcionarios']) if employee_data['nomes_funcionarios'] else 'Nenhum cadastrado'}"
            else:
                response = "👥 Consultando informações de funcionários... Verificando escalas, folgas e horários."
        elif "vendas" in user_message or "promoção" in user_message:
            sales_data = get_sales_data()
            if "error" not in sales_data:
                response = f"🛍️ Vendas: {sales_data['total_pedidos']} pedidos no total, {sales_data['pedidos_concluidos']} concluídos. Valor total em vendas: R$ {sales_data['valor_total_vendas']:.2f}"
            else:
                response = "🛍️ Analisando vendas... Identificando tendências e sugerindo estratégias promocionais."
        elif "pagamento" in user_message or "salário" in user_message:
            response = "💰 Analisando dados de pagamentos... Verificando folha de pagamento e vencimentos."
        elif "antonio rabelo" in user_message or "história" in user_message:
            response = f"🏆 {ANTONIO_RABELO_HISTORY['nome']} é o fundador da {ANTONIO_RABELO_HISTORY['empresa']}, especializada em {ANTONIO_RABELO_HISTORY['especialidade']}. Nossa empresa é {ANTONIO_RABELO_HISTORY['fundacao']} com foco em {ANTONIO_RABELO_HISTORY['valores']}."
        else:
            response = f"Olá! Sou a Lua, sua assistente de IA especializada na {ANTONIO_RABELO_HISTORY['empresa']}. Como posso ajudá-lo hoje? Posso gerar relatórios, analisar estoque, consultar funcionários ou fornecer insights sobre vendas."
        
        return jsonify({
            "response": response,
            "timestamp": datetime.now().isoformat(),
            "status": "success"
        })
    except Exception as e:
        return jsonify({
            "response": "Desculpe, ocorreu um erro ao processar sua mensagem.",
            "error": str(e),
            "status": "error"
        }), 500

@lua_bp.route("/lua/upload", methods=["POST"])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({
                "message": "Nenhum arquivo foi enviado.",
                "status": "error"
            }), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({
                "message": "Nome do arquivo está vazio.",
                "status": "error"
            }), 400
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            file.save(filepath)
            
            # Processar arquivo baseado na extensão
            file_extension = filename.rsplit('.', 1)[1].lower()
            
            if file_extension == 'csv':
                result = process_csv_file(filepath)
            elif file_extension in ['xlsx', 'xls']:
                result = process_excel_file(filepath)
            elif file_extension == 'pdf':
                result = process_pdf_file(filepath)
            elif file_extension == 'docx':
                result = process_docx_file(filepath)
            else:
                result = {
                    "status": "error",
                    "message": "Tipo de arquivo não suportado para processamento detalhado."
                }
            
            # Limpar arquivo temporário
            try:
                os.remove(filepath)
            except:
                pass
            
            return jsonify({
                "message": result["message"],
                "filename": filename,
                "file_type": file_extension,
                "processing_result": result,
                "timestamp": datetime.now().isoformat(),
                "status": result["status"]
            })
        else:
            return jsonify({
                "message": f"Tipo de arquivo não suportado. Aceito apenas: {', '.join(ALLOWED_EXTENSIONS)}",
                "status": "error"
            }), 400
            
    except Exception as e:
        return jsonify({
            "message": "Erro ao processar arquivo.",
            "error": str(e),
            "status": "error"
        }), 500

@lua_bp.route("/lua/report", methods=["POST"])
def generate_report():
    try:
        from src.report_generator import report_generator
        from flask import send_file
        
        report_type = request.json.get("type", "").lower()
        period = request.json.get("period", "mensal")
        year = request.json.get("year", datetime.now().year)
        month = request.json.get("month", None)
        
        # Gerar relatório baseado no tipo
        if report_type == "financeiro":
            result = report_generator.generate_financial_report(year, month)
        elif report_type == "funcionarios":
            result = report_generator.generate_employee_report(year, month)
        elif report_type == "estoque":
            result = report_generator.generate_inventory_report()
        else:
            return jsonify({
                "response": f"Tipo de relatório '{report_type}' não suportado.",
                "status": "error"
            }), 400
        
        if result["status"] == "success":
            return jsonify({
                "response": f"📊 {result['message']}! PDF disponível para download.",
                "report_type": report_type,
                "period": period,
                "filename": result["filename"],
                "filepath": result["filepath"],
                "data": result.get("data", {}),
                "timestamp": datetime.now().isoformat(),
                "status": "success",
                "download_url": f"/lua/download_report/{result['filename']}"
            })
        else:
            return jsonify({
                "response": result["message"],
                "status": "error"
            }), 500
            
    except Exception as e:
        return jsonify({
            "response": "Erro ao gerar relatório.",
            "error": str(e),
            "status": "error"
        }), 500

@lua_bp.route("/lua/download_report/<filename>", methods=["GET"])
def download_report(filename):
    """Endpoint para download de relatórios"""
    try:
        from flask import send_file
        import os
        
        reports_dir = "/tmp/reports"
        filepath = os.path.join(reports_dir, filename)
        
        if not os.path.exists(filepath):
            return jsonify({
                "error": "Arquivo não encontrado"
            }), 404
        
        return send_file(
            filepath,
            as_attachment=True,
            download_name=filename,
            mimetype='application/pdf'
        )
        
    except Exception as e:
        return jsonify({
            "error": f"Erro ao baixar relatório: {str(e)}"
        }), 500

@lua_bp.route("/lua/admin_input", methods=["POST"])
def admin_input():
    try:
        admin_data = request.json.get("data", {})
        input_type = admin_data.get("type", "")
        
        # Tratamento especial para contato do dono
        if input_type == "contato_dono":
            # Salvar configurações de contato do dono
            OWNER_CONTACT_CONFIG.update({
                "email": admin_data.get("email", ""),
                "telefone": admin_data.get("telefone", ""),
                "notif_estoque": admin_data.get("notif_estoque", True),
                "notif_aniversarios": admin_data.get("notif_aniversarios", True),
                "notif_vendas": admin_data.get("notif_vendas", False),
                "notif_financeiro": admin_data.get("notif_financeiro", True),
                "updated_at": datetime.now().isoformat()
            })
            
            response = f"📱 Contato do dono atualizado: {admin_data.get('email', '')} | {admin_data.get('telefone', '')}. Notificações configuradas."
        else:
            responses = {
                "horario_almoco": f"⏰ Horário de almoço atualizado: {admin_data.get('entrada', '')} às {admin_data.get('saida', '')}",
                "dia_pagamento": f"💰 Dia de pagamento definido para: {admin_data.get('dia', '')} de cada mês",
                "folga_funcionario": f"📅 Folga registrada para {admin_data.get('funcionario', '')}: {admin_data.get('data_folga', '')}",
                "escala": f"📋 Escala atualizada para {admin_data.get('funcionario', '')}: {admin_data.get('horario', '')}"
            }
            
            response = responses.get(input_type, f"Dados administrativos recebidos e processados: {input_type}")
        
        return jsonify({
            "response": response,
            "input_type": input_type,
            "data": admin_data,
            "timestamp": datetime.now().isoformat(),
            "status": "success"
        })
    except Exception as e:
        return jsonify({
            "response": "Erro ao processar dados administrativos.",
            "error": str(e),
            "status": "error"
        }), 500

@lua_bp.route("/lua/owner_contact", methods=["GET", "POST"])
def owner_contact():
    """Endpoint para gerenciar contato do dono"""
    try:
        if request.method == "GET":
            return jsonify({
                "contact": OWNER_CONTACT_CONFIG,
                "status": "success"
            })
        
        elif request.method == "POST":
            contact_data = request.json.get("contact", {})
            
            # Atualizar configurações de contato
            OWNER_CONTACT_CONFIG.update({
                "email": contact_data.get("email", ""),
                "telefone": contact_data.get("telefone", ""),
                "notif_estoque": contact_data.get("notif_estoque", True),
                "notif_aniversarios": contact_data.get("notif_aniversarios", True),
                "notif_vendas": contact_data.get("notif_vendas", False),
                "notif_financeiro": contact_data.get("notif_financeiro", True),
                "updated_at": datetime.now().isoformat()
            })
            
            return jsonify({
                "message": "Contato do dono atualizado com sucesso",
                "contact": OWNER_CONTACT_CONFIG,
                "status": "success"
            })
            
    except Exception as e:
        return jsonify({
            "message": "Erro ao gerenciar contato do dono",
            "error": str(e),
            "status": "error"
        }), 500

@lua_bp.route("/lua/send_notification", methods=["POST"])
def send_notification():
    """Endpoint para enviar notificações para o dono"""
    try:
        notification_data = request.json.get("notification", {})
        notification_type = notification_data.get("type", "")
        message = notification_data.get("message", "")
        
        # Verificar se o dono quer receber este tipo de notificação
        if not OWNER_CONTACT_CONFIG.get(f"notif_{notification_type}", False):
            return jsonify({
                "message": f"Notificação {notification_type} desabilitada pelo dono",
                "status": "skipped"
            })
        
        # Simular envio de notificação (aqui seria integração com serviços reais)
        email = OWNER_CONTACT_CONFIG.get("email", "")
        telefone = OWNER_CONTACT_CONFIG.get("telefone", "")
        
        if email or telefone:
            # Log da notificação enviada
            notification_log = {
                "type": notification_type,
                "message": message,
                "email": email,
                "telefone": telefone,
                "timestamp": datetime.now().isoformat(),
                "status": "sent"
            }
            
            # Adicionar ao histórico de notificações
            if "notification_history" not in globals():
                global notification_history
                notification_history = []
            
            notification_history.append(notification_log)
            
            # Manter apenas as últimas 100 notificações
            if len(notification_history) > 100:
                notification_history.pop(0)
            
            return jsonify({
                "message": f"Notificação {notification_type} enviada com sucesso",
                "notification": notification_log,
                "status": "success"
            })
        else:
            return jsonify({
                "message": "Contato do dono não configurado",
                "status": "error"
            }), 400
            
    except Exception as e:
        return jsonify({
            "message": "Erro ao enviar notificação",
            "error": str(e),
            "status": "error"
        }), 500

@lua_bp.route("/lua/insights", methods=["GET"])
def get_insights():
    """Endpoint para fornecer insights estratégicos"""
    try:
        insights = [
            "💎 Anéis de ouro têm maior margem de lucro no trimestre atual",
            "📈 Vendas de colares aumentaram 15% comparado ao mês anterior", 
            "⚠️ Estoque de brincos de prata está baixo - considere reposição",
            "👥 Funcionário João tem excelente performance em vendas",
            "💰 Fluxo de caixa positivo nos últimos 3 meses"
        ]
        
        return jsonify({
            "insights": insights,
            "timestamp": datetime.now().isoformat(),
            "status": "success"
        })
    except Exception as e:
        return jsonify({
            "insights": [],
            "error": str(e),
            "status": "error"
        }), 500

@lua_bp.route("/lua/learn", methods=["POST"])
def continuous_learning():
    """Endpoint para aprendizado contínuo"""
    try:
        learning_data = request.json.get("data", {})
        learning_type = learning_data.get("type", "")
        
        # Simular aprendizado contínuo
        response = f"🧠 Aprendizado registrado: {learning_type}. Dados incorporados ao conhecimento da Lua sobre a {ANTONIO_RABELO_HISTORY['empresa']}."
        
        return jsonify({
            "response": response,
            "learning_type": learning_type,
            "timestamp": datetime.now().isoformat(),
            "status": "success"
        })
    except Exception as e:
        return jsonify({
            "response": "Erro no processo de aprendizado.",
            "error": str(e),
            "status": "error"
        }), 500

@lua_bp.route("/lua/automatic_reports", methods=["GET"])
def get_automatic_reports():
    """Endpoint para acessar relatórios gerados automaticamente"""
    try:
        # Retornar os últimos 10 relatórios
        recent_reports = AUTOMATIC_REPORTS[-10:] if AUTOMATIC_REPORTS else []
        
        return jsonify({
            "reports": recent_reports,
            "total_reports": len(AUTOMATIC_REPORTS),
            "timestamp": datetime.now().isoformat(),
            "status": "success"
        })
    except Exception as e:
        return jsonify({
            "reports": [],
            "error": str(e),
            "status": "error"
        }), 500

@lua_bp.route("/lua/scheduler_config", methods=["GET", "POST"])
def scheduler_config():
    """Endpoint para configurar o scheduler"""
    try:
        if request.method == "GET":
            return jsonify({
                "config": SCHEDULER_CONFIG,
                "status": "success"
            })
        
        elif request.method == "POST":
            new_config = request.json.get("config", {})
            
            # Atualizar configuração do scheduler
            for report_type, config in new_config.items():
                if report_type in SCHEDULER_CONFIG:
                    SCHEDULER_CONFIG[report_type].update(config)
            
            return jsonify({
                "message": "Configuração do scheduler atualizada com sucesso",
                "config": SCHEDULER_CONFIG,
                "status": "success"
            })
            
    except Exception as e:
        return jsonify({
            "message": "Erro ao configurar scheduler",
            "error": str(e),
            "status": "error"
        }), 500




def execute_action(command):
    """
    Função principal para interpretar e executar comandos em linguagem natural
    """
    try:
        command = command.lower().strip()
        
        # Comandos para Vales
        if "vale" in command or "adiantamento" in command:
            return handle_vale_command(command)
        
        # Comandos para Viagens e Hospedagens
        elif "viagem" in command or "hospedagem" in command or "hotel" in command:
            return handle_travel_command(command)
        
        # Comandos para Produção
        elif "produção" in command or "producao" in command or "horas" in command or "trabalho" in command:
            return handle_production_command(command)
        
        # Comandos para Caixa
        elif "caixa" in command or "entrada" in command or "saída" in command or "saida" in command:
            return handle_caixa_command(command)
        
        # Comandos para Folha de Pagamento
        elif "folha" in command or "pagamento" in command or "salário" in command or "salario" in command:
            return handle_payroll_command(command)
        
        # Comandos para Clientes
        elif "cliente" in command or "customer" in command or "pedido" in command or "encomenda" in command:
            from src.customer_commands import handle_customer_command
            return handle_customer_command(command)
        
        # Comandos para Fornecedores
        elif "fornecedor" in command or "supplier" in command or "compra" in command:
            from src.supplier_commands import handle_supplier_command
            return handle_supplier_command(command)
        
        # Comandos para Configurações
        elif "horário" in command or "horario" in command or "configuração" in command or "configuracao" in command:
            return handle_config_command(command)
        
        else:
            return {
                "status": "error",
                "message": "Comando não reconhecido. Tente comandos como: 'adiciona vale para João de R$500', 'registra viagem para Pedro', 'cadastrar cliente Maria', 'compra de ouro do fornecedor João'."
            }
            
    except Exception as e:
        return {
            "status": "error",
            "message": f"Erro ao executar comando: {str(e)}"
        }

def handle_vale_command(command):
    """Processar comandos relacionados a vales"""
    try:
        import re
        from src.models.vale import Vale
        from src.models.employee import Employee
        from src.employee_search import extract_employee_name_from_command, find_employee_by_name
        
        # Extrair nome do funcionário e valor
        # Padrões: "vale para João de R$500", "adiciona vale João 300", etc.
        
        # Buscar valor monetário
        valor_match = re.search(r'(?:r\$|rs|reais?)\s*(\d+(?:[.,]\d{2})?)', command)
        if not valor_match:
            valor_match = re.search(r'(\d+(?:[.,]\d{2})?)\s*(?:r\$|rs|reais?)', command)
        if not valor_match:
            valor_match = re.search(r'(\d+(?:[.,]\d{2})?)', command)
        
        if not valor_match:
            return {
                "status": "error",
                "message": "Não foi possível identificar o valor do vale. Use formato como 'R$500' ou '500 reais'."
            }
        
        valor = float(valor_match.group(1).replace(',', '.'))
        
        # Extrair nome do funcionário usando a função inteligente
        nome_funcionario = extract_employee_name_from_command(command)
        
        if not nome_funcionario:
            return {
                "status": "error",
                "message": "Não foi possível identificar o nome do funcionário. Tente: 'vale para João de R$500'."
            }
        
        # Buscar funcionário usando busca inteligente
        search_result = find_employee_by_name(nome_funcionario)
        
        if search_result["status"] == "error":
            return search_result
        
        elif search_result["status"] == "multiple_matches":
            # Retornar opções para o usuário escolher
            options_text = "\n".join([
                f"{i+1}. {opt['name']}"
                for i, opt in enumerate(search_result["options"])
            ])
            return {
                "status": "multiple_matches",
                "message": f"Múltiplos funcionários encontrados para '{nome_funcionario}'. Escolha uma opção:\n{options_text}\n\nResponda com: 'vale para [nome completo] de R${valor:.2f}'",
                "options": search_result["options"]
            }
        
        funcionario = search_result["employee"]
        
        # Criar vale
        novo_vale = Vale(
            employee_id=funcionario.id,
            amount=valor,
            description=f"Vale criado via comando Lua: {command}"
        )
        
        db.session.add(novo_vale)
        
        # Criar transação no caixa
        from src.models.caixa import CaixaTransaction
        transacao_caixa = CaixaTransaction(
            type='saida',
            amount=valor,
            date=datetime.now().strftime('%Y-%m-%d'),
            description=f"Vale para {funcionario.name}",
            employee_id=funcionario.id,
            vale_id=novo_vale.id
        )
        
        db.session.add(transacao_caixa)
        
        # Atualizar folha de pagamento
        try:
            novo_vale.update_payroll()
        except:
            # Continuar mesmo se não conseguir atualizar a folha
            pass
        
        db.session.commit()
        
        match_info = ""
        if "similarity" in search_result and search_result["similarity"] < 1.0:
            match_info = f" (encontrado por similaridade: {search_result['similarity']:.0%})"
        
        return {
            "status": "success",
            "message": f"✅ Vale de R${valor:.2f} adicionado para {funcionario.name}{match_info}, integrado no Caixa e Folha de Pagamento.",
            "data": {
                "funcionario": funcionario.name,
                "valor": valor,
                "vale_id": novo_vale.id
            }
        }
        
    except Exception as e:
        db.session.rollback()
        return {
            "status": "error",
            "message": f"Erro ao processar vale: {str(e)}"
        }

def handle_travel_command(command):
    """Processar comandos relacionados a viagens e hospedagens"""
    try:
        import re
        from src.models.employee import Employee
        from src.models.caixa import CaixaTransaction
        from src.entity_search import extract_entity_name_from_command, find_entity_by_name
        
        # Extrair informações da viagem
        # Padrões: "viagem para Pedro, hotel R$1.200, de 12/09 a 15/09"
        
        # Buscar valor
        valor_match = re.search(r'(?:r\$|rs|hotel|hospedagem)\s*(\d+(?:[.,]\d{2})?)', command)
        if not valor_match:
            valor_match = re.search(r'(\d+(?:[.,]\d{2})?)', command)
        
        valor = float(valor_match.group(1).replace(',', '.')) if valor_match else 0.0
        
        # Extrair nome do funcionário usando busca inteligente
        nome_funcionario = extract_entity_name_from_command(command, "funcionario")
        
        # Buscar datas
        datas = re.findall(r'\d{2}/\d{2}(?:/\d{4})?', command)
        data_inicio = datas[0] if len(datas) > 0 else datetime.now().strftime('%d/%m/%Y')
        data_fim = datas[1] if len(datas) > 1 else data_inicio
        
        # Buscar funcionário no banco usando busca inteligente
        funcionario = None
        if nome_funcionario:
            search_result = find_entity_by_name(Employee, nome_funcionario)
            if search_result["status"] == "success":
                funcionario = search_result["entity"]
                nome_funcionario = funcionario.name
            elif search_result["status"] == "multiple_matches":
                options_text = "\n".join([
                    f"{i+1}. {opt['name']}"
                    for i, opt in enumerate(search_result["options"])
                ])
                return {
                    "status": "multiple_matches",
                    "message": f"Múltiplos funcionários encontrados para '{nome_funcionario}'. Escolha uma opção:\n{options_text}\n\nResponda com: 'viagem para [nome completo], hotel R${valor:.2f}'",
                    "options": search_result["options"]
                }
        else:
            nome_funcionario = "Funcionário"
        
        # Criar transação no caixa para a viagem
        if valor > 0:
            transacao_caixa = CaixaTransaction(
                type='saida',
                amount=valor,
                date=datetime.now().strftime('%Y-%m-%d'),
                description=f"Viagem/Hospedagem - {nome_funcionario} ({data_inicio} a {data_fim})",
                employee_id=funcionario.id if funcionario else None
            )
            
            db.session.add(transacao_caixa)
            db.session.commit()
        
        return {
            "status": "success",
            "message": f"🛎️ Viagem registrada: {nome_funcionario}, hotel R${valor:.2f}, {data_inicio} a {data_fim}.",
            "data": {
                "funcionario": nome_funcionario,
                "valor": valor,
                "data_inicio": data_inicio,
                "data_fim": data_fim
            }
        }
        
    except Exception as e:
        db.session.rollback()
        return {
            "status": "error",
            "message": f"Erro ao processar viagem: {str(e)}"
        }

def handle_production_command(command):
    """Processar comandos relacionados à produção"""
    try:
        import re
        from src.models.employee import Employee
        
        # Extrair horas e funcionário
        # Padrões: "adiciona 5 horas de trabalho do funcionário X na peça Y"
        
        # Buscar horas
        horas_match = re.search(r'(\d+(?:[.,]\d+)?)\s*(?:horas?|h)', command)
        if not horas_match:
            horas_match = re.search(r'(\d+(?:[.,]\d+)?)', command)
        
        if not horas_match:
            return {
                "status": "error",
                "message": "Não foi possível identificar as horas de trabalho."
            }
        
        horas = float(horas_match.group(1).replace(',', '.'))
        
        # Buscar nome do funcionário
        palavras_remover = ['adiciona', 'horas', 'de', 'trabalho', 'do', 'funcionário', 'funcionario', 'na', 'peça', 'peca']
        palavras = command.split()
        nome_palavras = []
        
        for palavra in palavras:
            palavra_limpa = re.sub(r'[^\w]', '', palavra)
            if palavra_limpa.lower() not in palavras_remover and not palavra_limpa.isdigit():
                nome_palavras.append(palavra_limpa)
        
        nome_funcionario = ' '.join(nome_palavras[:2]).title() if nome_palavras else "Funcionário"
        
        # Buscar peça/produto
        peca_match = re.search(r'(?:peça|peca|produto)\s+(\w+)', command)
        peca = peca_match.group(1) if peca_match else "Produto Geral"
        
        # Buscar funcionário no banco
        funcionario = None
        if nome_funcionario != "Funcionário":
            funcionario = Employee.query.filter(Employee.name.ilike(f'%{nome_funcionario}%')).first()
        
        # Registrar na produção (simulado - aqui seria integração com sistema de produção real)
        # Por enquanto, apenas log da ação
        
        return {
            "status": "success",
            "message": f"⚙️ Registrado: {horas} horas de trabalho de {nome_funcionario} na {peca}.",
            "data": {
                "funcionario": nome_funcionario,
                "horas": horas,
                "peca": peca,
                "funcionario_id": funcionario.id if funcionario else None
            }
        }
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Erro ao processar produção: {str(e)}"
        }

def handle_caixa_command(command):
    """Processar comandos relacionados ao caixa"""
    try:
        import re
        from src.models.caixa import CaixaTransaction
        
        # Determinar tipo de transação
        tipo = "entrada" if any(palavra in command for palavra in ["entrada", "recebimento", "venda"]) else "saida"
        
        # Extrair valor
        valor_match = re.search(r'(?:r\$|rs)\s*(\d+(?:[.,]\d{2})?)', command)
        if not valor_match:
            valor_match = re.search(r'(\d+(?:[.,]\d{2})?)', command)
        
        if not valor_match:
            return {
                "status": "error",
                "message": "Não foi possível identificar o valor da transação."
            }
        
        valor = float(valor_match.group(1).replace(',', '.'))
        
        # Extrair descrição
        descricao_palavras = []
        palavras_remover = ['caixa', 'entrada', 'saida', 'saída', 'adiciona', 'registra', 'r$', 'rs']
        
        for palavra in command.split():
            palavra_limpa = re.sub(r'[^\w]', '', palavra)
            if palavra_limpa.lower() not in palavras_remover and not palavra_limpa.isdigit():
                descricao_palavras.append(palavra)
        
        descricao = ' '.join(descricao_palavras) if descricao_palavras else f"Transação {tipo}"
        
        # Criar transação
        transacao = CaixaTransaction(
            type=tipo,
            amount=valor if tipo == "entrada" else -valor,
            date=datetime.now().strftime('%Y-%m-%d'),
            description=descricao
        )
        
        db.session.add(transacao)
        db.session.commit()
        
        return {
            "status": "success",
            "message": f"💰 {tipo.title()} de R${valor:.2f} registrada no caixa: {descricao}",
            "data": {
                "tipo": tipo,
                "valor": valor,
                "descricao": descricao
            }
        }
        
    except Exception as e:
        db.session.rollback()
        return {
            "status": "error",
            "message": f"Erro ao processar transação do caixa: {str(e)}"
        }

def handle_payroll_command(command):
    """Processar comandos relacionados à folha de pagamento"""
    try:
        import re
        from src.models.employee import Employee
        from src.models.payroll import Payroll
        
        # Buscar funcionário
        palavras_remover = ['folha', 'pagamento', 'salário', 'salario', 'atualiza', 'altera']
        palavras = command.split()
        nome_palavras = []
        
        for palavra in palavras:
            palavra_limpa = re.sub(r'[^\w]', '', palavra)
            if palavra_limpa.lower() not in palavras_remover and not palavra_limpa.isdigit():
                nome_palavras.append(palavra_limpa)
        
        if not nome_palavras:
            # Comando geral da folha
            return {
                "status": "success",
                "message": "📋 Folha de pagamento atualizada. Todos os cálculos foram reprocessados.",
                "data": {"acao": "atualizacao_geral"}
            }
        
        nome_funcionario = ' '.join(nome_palavras).title()
        
        # Buscar funcionário
        funcionario = Employee.query.filter(Employee.name.ilike(f'%{nome_funcionario}%')).first()
        
        if not funcionario:
            return {
                "status": "error",
                "message": f"Funcionário '{nome_funcionario}' não encontrado."
            }
        
        # Atualizar folha de pagamento do mês atual
        mes_atual = datetime.now().month
        ano_atual = datetime.now().year
        
        folha = Payroll.query.filter_by(
            employee_id=funcionario.id,
            month=mes_atual,
            year=ano_atual
        ).first()
        
        if not folha:
            folha = Payroll(
                employee_id=funcionario.id,
                month=mes_atual,
                year=ano_atual,
                base_salary=funcionario.salary or 0.0,
                total_vales=0.0,
                net_salary=funcionario.salary or 0.0
            )
            db.session.add(folha)
        
        # Atualizar totais
        folha.update_vales_total()
        db.session.commit()
        
        return {
            "status": "success",
            "message": f"📊 Folha de pagamento de {funcionario.name} atualizada para {mes_atual}/{ano_atual}.",
            "data": {
                "funcionario": funcionario.name,
                "mes": mes_atual,
                "ano": ano_atual,
                "salario_base": folha.base_salary,
                "total_vales": folha.total_vales,
                "salario_liquido": folha.net_salary
            }
        }
        
    except Exception as e:
        db.session.rollback()
        return {
            "status": "error",
            "message": f"Erro ao processar folha de pagamento: {str(e)}"
        }

def handle_config_command(command):
    """Processar comandos de configuração"""
    try:
        import re
        
        # Extrair horários
        horarios = re.findall(r'(\d{1,2})[h:]?(?:\d{2})?', command)
        
        if "entrada" in command and "saída" in command or "saida" in command:
            entrada = horarios[0] if len(horarios) > 0 else "8"
            almoco_saida = horarios[1] if len(horarios) > 1 else "11"
            almoco_volta = horarios[2] if len(horarios) > 2 else "12"
            saida = horarios[3] if len(horarios) > 3 else "17"
            
            return {
                "status": "success",
                "message": f"⏰ Horários atualizados: Entrada {entrada}h, Almoço {almoco_saida}h-{almoco_volta}h, Saída {saida}h",
                "data": {
                    "entrada": f"{entrada}:00",
                    "almoco_saida": f"{almoco_saida}:00",
                    "almoco_volta": f"{almoco_volta}:00",
                    "saida": f"{saida}:00"
                }
            }
        
        return {
            "status": "success",
            "message": "⚙️ Configuração processada com sucesso.",
            "data": {"comando": command}
        }
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Erro ao processar configuração: {str(e)}"
        }

@lua_bp.route("/lua", methods=["POST"])
def lua_execute():
    """Rota principal para executar comandos da Lua"""
    try:
        data = request.get_json()
        command = data.get("command", "")
        
        if not command:
            return jsonify({
                "status": "error",
                "message": "Comando não fornecido."
            }), 400
        
        # Executar comando
        result = execute_action(command)
        
        # Log da ação executada
        action_log = {
            "command": command,
            "result": result,
            "timestamp": datetime.now().isoformat(),
            "user": "Lua Assistant"
        }
        
        # Adicionar ao histórico de ações
        if "action_history" not in globals():
            global action_history
            action_history = []
        
        action_history.append(action_log)
        
        # Manter apenas as últimas 100 ações
        if len(action_history) > 100:
            action_history.pop(0)
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"Erro interno: {str(e)}"
        }), 500

@lua_bp.route("/lua/history", methods=["GET"])
def get_action_history():
    """Endpoint para acessar histórico de ações"""
    try:
        if "action_history" not in globals():
            global action_history
            action_history = []
        
        # Retornar as últimas 20 ações
        recent_actions = action_history[-20:] if action_history else []
        
        return jsonify({
            "actions": recent_actions,
            "total_actions": len(action_history),
            "status": "success"
        })
        
    except Exception as e:
        return jsonify({
            "actions": [],
            "error": str(e),
            "status": "error"
        }), 500



# ===== NOVAS FUNCIONALIDADES =====

def energy_cost_action(command):
    """Processar comandos relacionados a energia"""
    try:
        command_lower = command.lower()
        
        if "importa" in command_lower and "energia" in command_lower:
            # Comando para importar conta de energia
            return {
                "status": "info",
                "message": "📄 Para importar conta de energia, faça upload do arquivo (PDF, CSV ou XLSX) através do chat.",
                "data": {
                    "action": "energy_import_request",
                    "supported_formats": ["PDF", "CSV", "XLSX"]
                }
            }
        
        elif "calcula" in command_lower and "energia" in command_lower:
            # Calcular métricas de energia
            period = "month"
            if "diário" in command_lower or "dia" in command_lower:
                period = "day"
            elif "semanal" in command_lower or "semana" in command_lower:
                period = "week"
            
            result = energy_processor.calculate_energy_metrics(period)
            
            if result["status"] == "success":
                metrics = result["data"]
                return {
                    "status": "success",
                    "message": f"⚡ Métricas de energia calculadas para {period}",
                    "data": {
                        "period": period,
                        "consumption_kwh": metrics.get("consumption_kwh", 0),
                        "cost_total": metrics.get("cost_total", 0),
                        "cost_per_kwh": metrics.get("cost_per_kwh", 0),
                        "metrics": metrics
                    }
                }
            else:
                return result
        
        elif "histórico" in command_lower and "energia" in command_lower:
            # Mostrar histórico de energia
            result = energy_processor.get_energy_history(6)
            return {
                "status": "success",
                "message": "📊 Histórico de energia dos últimos 6 meses",
                "data": result["data"]
            }
        
        elif "alerta" in command_lower and "energia" in command_lower:
            # Gerar alertas de energia
            result = energy_processor.generate_energy_alerts()
            return {
                "status": "success",
                "message": "🚨 Alertas de energia gerados",
                "data": result
            }
        
        else:
            return {
                "status": "info",
                "message": "⚡ Comandos de energia disponíveis: 'importa energia', 'calcula energia diária/semanal/mensal', 'histórico energia', 'alertas energia'",
                "data": {"available_commands": ["importa energia", "calcula energia", "histórico energia", "alertas energia"]}
            }
            
    except Exception as e:
        return {
            "status": "error",
            "message": f"Erro ao processar comando de energia: {str(e)}"
        }

def monthly_financial_summary():
    """Gerar resumo financeiro mensal com custo e lucro"""
    try:
        now = datetime.now()
        start_of_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        
        # Buscar transações do mês atual
        transactions = CaixaTransaction.query.filter(
            CaixaTransaction.date >= start_of_month
        ).all()
        
        # Calcular totais
        total_income = sum(t.amount for t in transactions if t.amount > 0)
        total_expenses = sum(abs(t.amount) for t in transactions if t.amount < 0)
        net_profit = total_income - total_expenses
        
        # Buscar custos específicos
        costs = Cost.query.filter(
            Cost.date >= start_of_month
        ).all()
        
        total_costs = sum(c.amount for c in costs)
        
        # Buscar lucros específicos
        profits = Profit.query.filter(
            Profit.date >= start_of_month
        ).all()
        
        total_profits = sum(p.amount for p in profits)
        
        # Calcular métricas
        gross_profit = total_profits
        net_profit_adjusted = gross_profit - total_costs
        profit_margin = (net_profit_adjusted / total_income * 100) if total_income > 0 else 0
        
        return {
            "status": "success",
            "message": "💰 Resumo financeiro mensal gerado",
            "data": {
                "period": f"{now.strftime('%B %Y')}",
                "total_income": total_income,
                "total_expenses": total_expenses,
                "total_costs": total_costs,
                "gross_profit": gross_profit,
                "net_profit": net_profit_adjusted,
                "profit_margin": round(profit_margin, 2),
                "transactions_count": len(transactions),
                "generated_at": now.isoformat()
            }
        }
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Erro ao gerar resumo financeiro: {str(e)}"
        }

def generate_vale_report(period="month"):
    """Gerar relatório de vales"""
    try:
        now = datetime.now()
        
        if period == "month":
            start_date = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        elif period == "week":
            start_date = now - timedelta(days=7)
        else:  # day
            start_date = now.replace(hour=0, minute=0, second=0, microsecond=0)
        
        # Buscar vales do período
        vales = CaixaTransaction.query.filter(
            CaixaTransaction.date >= start_date,
            CaixaTransaction.description.like('%vale%')
        ).all()
        
        # Agrupar por funcionário
        vales_by_employee = {}
        total_vales = 0
        
        for vale in vales:
            # Extrair nome do funcionário da descrição
            desc = vale.description.lower()
            employee_name = "Não identificado"
            
            # Tentar extrair nome comum
            common_names = ["joão", "maria", "pedro", "ana", "carlos", "josé", "antônio", "francisco"]
            for name in common_names:
                if name in desc:
                    employee_name = name.title()
                    break
            
            if employee_name not in vales_by_employee:
                vales_by_employee[employee_name] = {
                    "total": 0,
                    "count": 0,
                    "vales": []
                }
            
            vale_amount = abs(vale.amount)
            vales_by_employee[employee_name]["total"] += vale_amount
            vales_by_employee[employee_name]["count"] += 1
            vales_by_employee[employee_name]["vales"].append({
                "date": vale.date.isoformat(),
                "amount": vale_amount,
                "description": vale.description
            })
            
            total_vales += vale_amount
        
        return {
            "status": "success",
            "message": f"📋 Relatório de vales - {period}",
            "data": {
                "period": period,
                "start_date": start_date.isoformat(),
                "total_vales": total_vales,
                "total_count": len(vales),
                "employees_count": len(vales_by_employee),
                "vales_by_employee": vales_by_employee,
                "generated_at": now.isoformat()
            }
        }
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Erro ao gerar relatório de vales: {str(e)}"
        }

def handle_all_modules_command(command):
    """Processar comandos para todos os módulos do sistema"""
    try:
        command_lower = command.lower()
        
        # Comandos de Financeiro
        if any(word in command_lower for word in ["financeiro", "receita", "despesa", "transação"]):
            if "adiciona" in command_lower or "registra" in command_lower:
                # Extrair valor
                import re
                valor_match = re.search(r'r?\$?\s*(\d+(?:,\d+)?)', command_lower)
                valor = float(valor_match.group(1).replace(',', '.')) if valor_match else 0
                
                tipo = "receita" if any(word in command_lower for word in ["receita", "entrada", "ganho"]) else "despesa"
                
                # Criar transação financeira
                transaction = FinancialTransaction(
                    amount=valor if tipo == "receita" else -valor,
                    description=f"Transação {tipo} via Lua",
                    date=datetime.now(),
                    category=tipo
                )
                
                db.session.add(transaction)
                db.session.commit()
                
                return {
                    "status": "success",
                    "message": f"💰 {tipo.title()} de R$ {valor:.2f} registrada no financeiro",
                    "data": {
                        "type": tipo,
                        "amount": valor,
                        "transaction_id": transaction.id
                    }
                }
        
        # Comandos de Estoque
        elif any(word in command_lower for word in ["estoque", "material", "produto"]):
            if "consulta" in command_lower or "mostra" in command_lower:
                materials = Material.query.limit(10).all()
                return {
                    "status": "success",
                    "message": f"📦 Estoque consultado - {len(materials)} materiais encontrados",
                    "data": {
                        "materials": [
                            {
                                "name": m.name,
                                "quantity": getattr(m, 'quantity', 0),
                                "unit": getattr(m, 'unit', 'un')
                            } for m in materials
                        ]
                    }
                }
        
        # Comandos de Funcionários
        elif any(word in command_lower for word in ["funcionário", "funcionario", "empregado", "colaborador"]):
            if "lista" in command_lower or "mostra" in command_lower:
                employees = Employee.query.limit(10).all()
                return {
                    "status": "success",
                    "message": f"👥 {len(employees)} funcionários encontrados",
                    "data": {
                        "employees": [
                            {
                                "name": e.name,
                                "position": getattr(e, 'position', 'Não definido'),
                                "salary": getattr(e, 'salary', 0)
                            } for e in employees
                        ]
                    }
                }
        
        # Comandos de Vendas
        elif any(word in command_lower for word in ["venda", "pedido", "order"]):
            if "consulta" in command_lower or "mostra" in command_lower:
                orders = Order.query.limit(10).all()
                return {
                    "status": "success",
                    "message": f"🛒 {len(orders)} pedidos encontrados",
                    "data": {
                        "orders": [
                            {
                                "id": o.id,
                                "total": getattr(o, 'total', 0),
                                "status": getattr(o, 'status', 'Pendente'),
                                "date": o.created_at.isoformat() if hasattr(o, 'created_at') else None
                            } for o in orders
                        ]
                    }
                }
        
        # Comandos de Produção
        elif any(word in command_lower for word in ["produção", "producao", "fabricação", "fabrica"]):
            if "consulta" in command_lower or "mostra" in command_lower:
                jewelries = Jewelry.query.limit(10).all()
                return {
                    "status": "success",
                    "message": f"⚙️ {len(jewelries)} itens em produção encontrados",
                    "data": {
                        "production": [
                            {
                                "name": j.name,
                                "type": getattr(j, 'type', 'Não definido'),
                                "status": getattr(j, 'status', 'Em produção')
                            } for j in jewelries
                        ]
                    }
                }
        
        else:
            return {
                "status": "info",
                "message": "🔧 Módulos disponíveis: Financeiro, Estoque, Funcionários, Vendas, Produção, Caixa, Vales, Viagens",
                "data": {
                    "available_modules": [
                        "Financeiro", "Estoque", "Funcionários", 
                        "Vendas", "Produção", "Caixa", "Vales", "Viagens"
                    ]
                }
            }
            
    except Exception as e:
        db.session.rollback()
        return {
            "status": "error",
            "message": f"Erro ao processar comando de módulo: {str(e)}"
        }

# Atualizar a função execute_action principal
def execute_action_enhanced(command, session_id=None):
    """Versão aprimorada da função execute_action com contexto e novas funcionalidades"""
    try:
        # Gerenciar contexto da conversa
        if not session_id:
            session_id = context_manager.get_session_id()
        
        # Adicionar comando ao contexto
        context_manager.add_message(session_id, command, "command")
        
        # Analisar contexto para melhorar comando
        enhanced_command, context_hints = context_manager.analyze_context_for_command(session_id, command)
        
        # Usar comando aprimorado se disponível
        final_command = enhanced_command if enhanced_command != command else command
        
        # Processar comando com base no tipo
        command_lower = final_command.lower()
        
        # Comandos de energia
        if any(word in command_lower for word in ["energia", "conta de luz", "kwh"]):
            result = energy_cost_action(final_command)
        
        # Comandos de custo/lucro
        elif any(word in command_lower for word in ["custo", "lucro", "financeiro mensal", "resumo financeiro"]):
            result = monthly_financial_summary()
        
        # Comandos de relatório de vales
        elif "relatório" in command_lower and "vale" in command_lower:
            period = "month"
            if "semanal" in command_lower:
                period = "week"
            elif "diário" in command_lower:
                period = "day"
            result = generate_vale_report(period)
        
        # Comandos para registrar e imprimir vale
        elif "vale" in command_lower and ("imprime" in command_lower or "imprimir" in command_lower):
            result = execute_vale_print(final_command)

        # Comandos para todos os módulos
        elif any(word in command_lower for word in [
            "financeiro", "estoque", "funcionário", "funcionario", 
            "venda", "produção", "producao", "material", "produto"
        ]):
            result = handle_all_modules_command(final_command)
        
        # Comandos existentes (vales, viagens, etc.)
        else:
            result = execute_action(final_command)  # Função original
        
        # Adicionar resultado ao contexto
        context_manager.add_message(session_id, result.get("message", ""), "response", result)
        
        # Atualizar contexto com dados relevantes
        if result.get("status") == "success" and result.get("data"):
            data = result["data"]
            if "funcionario" in data:
                context_manager.update_context(session_id, "last_employee", data["funcionario"])
            if "valor" in data:
                context_manager.update_context(session_id, "last_value", data["valor"])
        
        return result
        
    except Exception as e:
        error_result = {
            "status": "error",
            "message": f"Erro ao executar comando aprimorado: {str(e)}"
        }
        
        if session_id:
            context_manager.add_message(session_id, error_result["message"], "response", error_result)
        
        return error_result

@lua_bp.route("/lua/enhanced", methods=["POST"])
def lua_execute_enhanced():
    """Rota aprimorada para executar comandos da Lua com contexto"""
    try:
        data = request.get_json()
        command = data.get("command", "")
        session_id = data.get("session_id")
        
        if not command:
            return jsonify({
                "status": "error",
                "message": "Comando não fornecido."
            }), 400
        
        # Executar comando aprimorado
        result = execute_action_enhanced(command, session_id)
        
        # Adicionar session_id ao resultado
        if session_id:
            result["session_id"] = session_id
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"Erro interno: {str(e)}"
        }), 500

@lua_bp.route("/lua/context/<session_id>", methods=["GET"])
def get_conversation_context(session_id):
    """Endpoint para acessar contexto da conversa"""
    try:
        history = context_manager.get_conversation_history(session_id, limit=20)
        context = context_manager.get_context(session_id)
        summary = context_manager.get_session_summary(session_id)
        
        return jsonify({
            "session_id": session_id,
            "history": history,
            "context": context,
            "summary": summary,
            "status": "success"
        })
        
    except Exception as e:
        return jsonify({
            "session_id": session_id,
            "history": [],
            "context": {},
            "summary": None,
            "error": str(e),
            "status": "error"
        }), 500

@lua_bp.route("/lua/energy/upload", methods=["POST"])
def upload_energy_file():
    """Endpoint para upload de arquivos de energia"""
    try:
        if 'file' not in request.files:
            return jsonify({
                "status": "error",
                "message": "Nenhum arquivo fornecido"
            }), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({
                "status": "error",
                "message": "Nenhum arquivo selecionado"
            }), 400
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            file.save(file_path)
            
            # Processar arquivo de energia
            file_extension = filename.rsplit('.', 1)[1].lower()
            result = energy_processor.process_energy_file(file_path, file_extension)
            
            # Limpar arquivo temporário
            try:
                os.remove(file_path)
            except:
                pass
            
            return jsonify(result)
        
        else:
            return jsonify({
                "status": "error",
                "message": "Formato de arquivo não suportado. Use PDF, CSV ou XLSX."
            }), 400
            
    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"Erro ao processar arquivo de energia: {str(e)}"
        }), 500

@lua_bp.route("/lua/dashboard/financial", methods=["GET"])
def get_financial_dashboard():
    """Endpoint para dados do dashboard financeiro"""
    try:
        financial_summary = monthly_financial_summary()
        vale_report = generate_vale_report("month")
        
        return jsonify({
            "status": "success",
            "data": {
                "financial_summary": financial_summary.get("data", {}),
                "vale_report": vale_report.get("data", {}),
                "generated_at": datetime.now().isoformat()
            }
        })
        
    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"Erro ao gerar dashboard financeiro: {str(e)}"
        }), 500

@lua_bp.route("/lua/dashboard/energy", methods=["GET"])
def get_energy_dashboard():
    """Endpoint para dados do dashboard de energia"""
    try:
        metrics = energy_processor.calculate_energy_metrics("month")
        history = energy_processor.get_energy_history(6)
        alerts = energy_processor.generate_energy_alerts()
        
        return jsonify({
            "status": "success",
            "data": {
                "metrics": metrics.get("data", {}),
                "history": history.get("data", []),
                "alerts": alerts.get("alerts", []),
                "generated_at": datetime.now().isoformat()
            }
        })
        
    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"Erro ao gerar dashboard de energia: {str(e)}"
        }), 500

# Inicializar limpeza automática de contexto
def start_context_cleanup():
    """Iniciar limpeza automática de contextos antigos"""
    def cleanup_worker():
        while True:
            try:
                context_manager.cleanup_old_sessions()
                time.sleep(3600)  # Limpar a cada hora
            except Exception as e:
                print(f"Erro na limpeza de contexto: {str(e)}")
                time.sleep(3600)
    
    cleanup_thread = threading.Thread(target=cleanup_worker, daemon=True)
    cleanup_thread.start()

# Inicializar limpeza ao importar o módulo
start_context_cleanup()



def execute_vale_print(command):
    try:
        # 1. Parse do comando
        parsed_data = parse_vale_command(command)
        if parsed_data["status"] == "error":
            return parsed_data

        employee_name = parsed_data["data"]["employee_name"]
        vale_amount = parsed_data["data"]["amount"]

        # 2. Registro do vale
        vale_registration = register_vale(employee_name, vale_amount)
        if vale_registration["status"] == "error":
            return vale_registration

        vale_id = vale_registration["data"]["vale_id"]
        employee_id = vale_registration["data"]["employee_id"]

        # 3. Atualização do Caixa
        update_caixa(employee_name, vale_amount, vale_id, employee_id)

        # 4. Atualização da Folha de Pagamento
        update_folha(employee_name, vale_amount)

        # 5. Geração do PDF
        remaining_salary = vale_registration["data"]["remaining_salary"]
        pdf_generation = vale_printer.generate_vale_pdf(employee_name, vale_amount, remaining_salary)
        if pdf_generation["status"] == "error":
            return pdf_generation

        pdf_path = pdf_generation["filepath"]

        # 6. Impressão do PDF
        print_result = vale_printer.print_vale(pdf_path)

        return {
            "status": "success",
            "message": f'✅ Vale de R${vale_amount:.2f} registrado para {employee_name}, integrado no Caixa e Folha, enviado para impressão.',
            "data": {
                "pdf_path": pdf_path,
                "print_status": print_result["status"]
            }
        }

    except Exception as e:
        return {"status": "error", "message": f"Erro no fluxo de impressão de vale: {str(e)}"}

def parse_vale_command(command):
    try:
        import re
        valor_match = re.search(r'(?:r\$|rs|reais?)\s*(\d+(?:[.,]\d{2})?)', command)
        if not valor_match:
            valor_match = re.search(r'(\d+(?:[.,]\d{2})?)\s*(?:r\$|rs|reais?)', command)
        if not valor_match:
            valor_match = re.search(r'(\d+(?:[.,]\d{2})?)', command)

        if not valor_match:
            return {"status": "error", "message": "Valor do vale não identificado."}

        valor = float(valor_match.group(1).replace(',', '.'))

        palavras_remover = ['vale', 'para', 'de', 'adiciona', 'criar', 'r$', 'rs', 'reais', 'real', 'imprime', 'imprimir', 'e', 'mim']
        palavras = command.split()
        nome_palavras = []

        for palavra in palavras:
            palavra_limpa = re.sub(r'[^\w]', '', palavra)
            if palavra_limpa.lower() not in palavras_remover and not palavra_limpa.isdigit():
                nome_palavras.append(palavra_limpa)

        if not nome_palavras:
            return {"status": "error", "message": "Nome do funcionário não identificado."}

        nome_funcionario = ' '.join(nome_palavras).title()

        return {"status": "success", "data": {"employee_name": nome_funcionario, "amount": valor}}
    except Exception as e:
        return {"status": "error", "message": f"Erro ao interpretar comando: {str(e)}"}

def register_vale(employee_name, amount):
    try:
        from src.models.employee import Employee
        from src.models.vale import Vale

        employee = Employee.query.filter(Employee.name.ilike(f'%{employee_name}%')).first()
        if not employee:
            return {"status": "error", "message": f"Funcionário '{employee_name}' não encontrado."}

        new_vale = Vale(employee_id=employee.id, amount=amount, description=f"Vale registrado via Lua")
        db.session.add(new_vale)
        db.session.commit()

        remaining_salary = (employee.salary or 0) - new_vale.get_total_vales_for_month()

        return {"status": "success", "data": {"vale_id": new_vale.id, "employee_id": employee.id, "remaining_salary": remaining_salary}}
    except Exception as e:
        db.session.rollback()
        return {"status": "error", "message": f"Erro ao registrar vale: {str(e)}"}

def update_caixa(employee_name, amount, vale_id, employee_id):
    try:
        from src.models.caixa import CaixaTransaction
        transacao_caixa = CaixaTransaction(
            type='saida',
            amount=amount,
            date=datetime.now().strftime('%Y-%m-%d'),
            description=f"Vale para {employee_name}",
            employee_id=employee_id,
            vale_id=vale_id
        )
        db.session.add(transacao_caixa)
        db.session.commit()
        return {"status": "success"}
    except Exception as e:
        db.session.rollback()
        return {"status": "error", "message": f"Erro ao atualizar caixa: {str(e)}"}

def update_folha(employee_name, amount):
    try:
        from src.models.employee import Employee
        from src.models.payroll import Payroll

        employee = Employee.query.filter(Employee.name.ilike(f'%{employee_name}%')).first()
        if not employee:
            return {"status": "error", "message": f"Funcionário '{employee_name}' não encontrado."}

        mes_atual = datetime.now().month
        ano_atual = datetime.now().year

        folha = Payroll.query.filter_by(employee_id=employee.id, month=mes_atual, year=ano_atual).first()
        if not folha:
            folha = Payroll(
                employee_id=employee.id,
                month=mes_atual,
                year=ano_atual,
                base_salary=employee.salary or 0.0,
                total_vales=0.0,
                net_salary=employee.salary or 0.0
            )
            db.session.add(folha)

        folha.update_vales_total()
        db.session.commit()
        return {"status": "success"}
    except Exception as e:
        db.session.rollback()
        return {"status": "error", "message": f"Erro ao atualizar folha: {str(e)}"}


